from http.client import HTTPException
from fastapi import APIRouter, Depends, Form
from sqlalchemy.orm import Session
import json
from docx import Document
import os
from fastapi.responses import FileResponse
from starlette.background import BackgroundTask
from FAST_scriber_app.config import TEXT_DIR
from FAST_scriber_app.database import get_db
from FAST_scriber_app.models import AudioFile

router = APIRouter()



def remove_file(path: str):
    try:
        os.remove(path)
    except Exception as e:
        print(f"Ошибка при удалении файла: {e}")


from fastapi import HTTPException

@router.post("/download/{audio_id}", summary="Скачивание файлов")
async def download_transcript(audio_id: str, format: str = Form('docx'), db: Session = Depends(get_db)):
    try:
        audio = db.get(AudioFile, audio_id)
        if not audio or not audio.transcript_uuid_name:
            raise HTTPException(status_code=404, detail="Аудиофайл или транскрипт не найден")

        text_path = TEXT_DIR / audio.transcript_uuid_name
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Файл транскрипта отсутствует")

        with open(text_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            segments = data.get('text', {}).get('segments', [])
            texts = [s.get('text', '').strip() for s in segments if s.get('text')]
            full_text = "\n".join(texts)

        filename = f"{audio.original_filename.rsplit('.', 1)[0]}.{format}"

        if format == "txt":
            with open(filename, 'w', encoding="utf-8") as file:
                file.write(full_text)
            return FileResponse(
                path=filename,
                background=BackgroundTask(remove_file, path=filename),
                media_type='text/plain',
                filename=filename
            )

        elif format == "docx":
            doc = Document()
            doc.add_heading(f'Транскрипт {audio.original_filename}', level=1)
            doc.add_paragraph(full_text)
            doc.save(filename)
            return FileResponse(
                path=filename,
                background=BackgroundTask(remove_file, path=filename),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                filename=filename
            )

        else:
            raise HTTPException(status_code=400, detail="Неподдерживаемый формат")

    except Exception as e:
        print(f"Ошибка при скачивании: {e}")
        raise HTTPException(status_code=500, detail="Ошибка при обработке запроса")

